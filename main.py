import os
import threading
import re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from PIL import Image
import tkinter as tk
from tkinter import filedialog
import fitz

output_file_name = "pretenzija.docx"
output_file_path = f"C:/Users/josta/Downloads/{output_file_name}"
pdf = None
output_doc = Document()
header_image_path = "CM_logo.png"
page_bottom_limit = 100


def register_fonts():
    pdfmetrics.registerFont(TTFont('Arial', 'Arial-Unicode-Regular.ttf'))
    pdfmetrics.registerFont(TTFont('Arial-Bold', 'Arial-Unicode-Bold.ttf'))
    pdfmetrics.registerFont(TTFont('Arial-Bold-Italic', 'Arial-Unicode-Bold-Italic.ttf'))
    pdfmetrics.registerFont(TTFont('Arial-Italic', 'Arial-Unicode-Italic.ttf'))


def add_text_to_docx(input_doc):
    global output_doc

    for para in input_doc.paragraphs:
        new_para = output_doc.add_paragraph()

        new_para.style = para.style

        for run in para.runs:
            new_run = new_para.add_run(run.text)

            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline

            if run.font.size:
                new_run.font.size = run.font.size

            if run.font.name:
                new_run.font.name = run.font.name

        if para.alignment is not None:
            new_para.alignment = para.alignment


def select_input_file():
    global output_file_path
    file_path = filedialog.askopenfilename()

    if not file_path:
        docx_file_label.config(text=f"Error opening file {file_path}")
        loading_label.update_idletasks()
        return

    docx_file_label.config(text="Word document selected")
    loading_label.update_idletasks()

    input_document = Document(file_path)
    add_text_to_docx(input_document)


def remove_table_borders(table):
    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)


def add_header_footer():
    global output_doc
    global header_image_path

    for section in output_doc.sections:

        header = section.header
        header_para = header.add_paragraph()

        try:
            header_para.add_run().add_picture(header_image_path, width=Pt(100))
        except Exception as e:
            print(f"Error loading header image: {e}")

        footer = section.footer
        footer_table = footer.add_table(rows=1, cols=2, width=Pt(500))  # Create a table with 1 row and 2 columns

        remove_table_borders(footer_table)

        footer_table.columns[0].width = Pt(250)
        footer_table.columns[1].width = Pt(250)

        cell_1 = footer_table.cell(0, 0)
        cell_1.text = "Kauno g. 16-308, LT-03212 Vilnius, Lietuva\n" \
                      "Įm. k. 305594385 \n" \
                      "UAB „Claims management“"

        cell_2 = footer_table.cell(0, 1)
        cell_2.text = "www.claimsmanagement.lt\n" \
                      "El. p.: paulius@claimsmanagement.lt\n" \
                      "Tel.nr.: +370 6 877 63 30"


def extract_images_from_pdf(pdf_path, output_folder):
    pdf_document = fitz.open(pdf_path)
    images = []
    image_xrefs = set()

    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        for img in page.get_images(full=True):
            xref = img[0]  # The image reference (xref)

            # Check if this image has already been extracted
            if xref in image_xrefs:
                continue  # Skip this image since it's already processed

            image_xrefs.add(xref)  # Mark the xref as processed
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image_name = f"pdf_image {page_num} {xref}.{image_ext}"

            # Save the extracted image temporarily to the output folder
            image_path = os.path.join(output_folder, image_name)
            with open(image_path, "wb") as f:
                f.write(image_bytes)

            images.append(image_name)

    pdf_document.close()
    return images


def paste_images_to_word_2x2(image_files, subfolder_path, num_of_images):
    global output_doc
    image_count = 0
    table = output_doc.add_table(rows=2, cols=2)

    for image_file in image_files:
        image_path = os.path.join(subfolder_path, image_file)

        try:
            img = Image.open(image_path)
            img = img.convert("RGB")
            img_width, img_height = img.size

            max_width = 200
            max_height = 260
            ratio = min(max_width / img_width, max_height / img_height)
            new_size = (int(img_width * ratio), int(img_height * ratio))

            row = (image_count % 4) // 2
            col = (image_count % 4) % 2

            cell = table.cell(row, col)
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph.add_run().add_picture(image_path, width=Pt(new_size[0]))

        except Exception as e:
            print(f"Error processing image {image_file}, num_of_images={num_of_images}: {e}")
            continue

        image_count += 1

        if image_count % 4 == 0 and image_count != num_of_images:
            output_doc.add_page_break()
            table = output_doc.add_table(rows=2, cols=2)


def paste_images_to_word_1pic(image_files, subfolder_path):
    global output_doc

    for image_file in image_files:
        image_path = os.path.join(subfolder_path, image_file)

        try:
            img = Image.open(image_path)
            img = img.convert("RGB")
            img_width, img_height = img.size

            max_width = 400
            max_height = 520
            ratio = min(max_width / img_width, max_height / img_height)
            new_size = (int(img_width * ratio), int(img_height * ratio))

            output_doc.add_picture(image_path, width=Pt(new_size[0]))

        except Exception as e:
            print(f"Error processing image {image_file}: {e}")
            continue


def process_images():
    global output_doc

    # Select the main folder
    folder_path = filedialog.askdirectory()
    if not folder_path:
        loading_label.config(text="No folder selected!")
        return

    subfolders = [f for f in os.listdir(folder_path) if os.path.isdir(os.path.join(folder_path, f))]

    # Sort subfolders by the number at the beginning of the folder name
    def extract_number(folder_name):
        match = re.match(r"(\d+)", folder_name)
        return int(match.group(0)) if match else float('inf')  # Sort non-matching folders last

    def extract_folder_free_text(folder_name):
        match = re.match(r"\d+ (.+)", folder_name)
        return match.group(1) if match else ""

    sorted_subfolders = sorted(subfolders, key=extract_number)

    for subfolder_name in sorted_subfolders:
        subfolder_path = os.path.join(folder_path, subfolder_name)

        if os.path.isdir(subfolder_path):
            # Add a new section or page for each subfolder
            output_doc.add_paragraph(f"{extract_folder_free_text(subfolder_name)}", style='Heading 1')

            image_files = [f for f in os.listdir(subfolder_path)
                           if f.lower().endswith(('.png', '.jpg', '.jpeg', '.webp', '.heic', '.pdf'))]

            # Extract images from PDFs if any
            for image_file in image_files:
                if image_file.lower().endswith('.pdf'):
                    pdf_path = os.path.join(subfolder_path, image_file)
                    extracted_images = extract_images_from_pdf(pdf_path, subfolder_path)
                    image_files.extend(extracted_images)

            print(image_files)
            # Remove the pdf files from the list after extracting images
            image_files = [f for f in image_files if not f.lower().endswith('.pdf')]
            # Remove duplicates from the list
            image_files = list(dict.fromkeys(image_files))
            print(image_files)

            if not image_files:
                print(f"No images found in folder: {subfolder_name}")
                continue

            num_of_images = len(image_files)

            if num_of_images > 1:
                paste_images_to_word_2x2(image_files, subfolder_path, num_of_images)
            elif num_of_images == 1:
                paste_images_to_word_1pic(image_files, subfolder_path)

    btn_generate_pdf.pack(pady=0)
    save_status_label.pack(pady=2)
    loading_label.config(text="Images processed successfully!")
    loading_label.update_idletasks()


def select_image_folder():
    loading_label.config(text="Processing images, please wait...")
    loading_label.update_idletasks()

    thread = threading.Thread(target=process_images)
    thread.start()


def select_output_folder():
    global output_file_path
    folder_path = filedialog.askdirectory()

    if folder_path:
        output_file_path = f"{folder_path}/{output_file_name}"
        output_folder_label.config(text=f"Output folder selected: {folder_path}")
        loading_label.update_idletasks()
    else:
        output_folder_label.config(text="Error selecting output folder")
        loading_label.update_idletasks()


def save_word():
    global output_file_path

    add_header_footer()
    output_doc.save(output_file_path)
    save_status_label.config(text="PDF saved")
    loading_label.update_idletasks()


if __name__ == '__main__':
    root = tk.Tk()
    root.title("CM PDF Converter")

    # register_fonts()

    btn_select_output_folder = tk.Button(root, text="Select Output Folder", command=select_output_folder)
    btn_select_docx = tk.Button(root, text="Select Input Word Document", command=select_input_file)
    btn_select_image_folder = tk.Button(root, text="Select Input Image Folder", command=select_image_folder)
    btn_generate_pdf = tk.Button(root, text=f"Generate {output_file_name}", command=save_word)

    output_folder_label = tk.Label(root, text="")
    docx_file_label = tk.Label(root, text="")
    loading_label = tk.Label(root, text="")
    save_status_label = tk.Label(root, text="")

    btn_select_output_folder.pack(pady=10)
    output_folder_label.pack(pady=0)
    btn_select_docx.pack(pady=10)
    docx_file_label.pack(pady=2)
    btn_select_image_folder.pack(pady=10)
    loading_label.pack(pady=2)
    btn_generate_pdf.forget()
    save_status_label.forget()

    root.mainloop()
